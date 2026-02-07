"""
Word文档生成器 - AI英语教学系统

使用 python-docx 库将教案数据转换为格式化的 Word 文档。
支持中文内容、表格、样式设置等功能。
"""
import logging
from io import BytesIO
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


class WordDocumentGenerator:
    """
    Word 文档生成器

    将教案数据转换为格式化的 Word (.docx) 文档。

    功能：
    - 生成包含教案完整内容的 Word 文档
    - 支持中文内容（自动设置中文字体）
    - 使用表格展示词汇和语法点
    - 使用编号列表展示练习题
    - 支持自定义样式和格式

    使用示例：
        ```python
        generator = WordDocumentGenerator()

        content = {
            "title": "过去完成时教学",
            "level": "B1",
            "topic": "Grammar",
            "duration": 45,
            # ... 其他教案内容
        }

        template_vars = {
            "teacher_name": "张老师",
            "school": "XX中学",
            "date": "2026-02-07"
        }

        doc_bytes = generator.generate(content, template_vars)
        ```
    """

    # 样式配置
    # 标题字体大小
    FONT_SIZE_TITLE = 18
    FONT_SIZE_HEADING1 = 14
    FONT_SIZE_HEADING2 = 12
    FONT_SIZE_NORMAL = 10.5

    # 行距
    LINEspacing_SINGLE = 1.0
    LINEspacing_1_5 = 1.5
    LINEspacing_DOUBLE = 2.0

    # 颜色配置
    COLOR_HEADING = RGBColor(0, 51, 102)  # 深蓝色
    COLOR_ACCENT = RGBColor(204, 0, 0)    # 红色用于强调
    COLOR_NORMAL = RGBColor(51, 51, 51)   # 深灰色

    def __init__(self):
        """初始化 Word 文档生成器"""
        self.doc: Optional[Document] = None

    def generate(self, content: Dict[str, Any], template_vars: Dict[str, Any]) -> bytes:
        """
        生成完整的 Word 文档

        Args:
            content: 教案内容数据，包含以下字段：
                - title: 教案标题
                - level: CEFR等级
                - topic: 主题
                - duration: 课程时长（分钟）
                - target_exam: 目标考试（可选）
                - objectives: 教学目标（字典）
                - vocabulary: 词汇表（字典）
                - grammar_points: 语法点（列表）
                - teaching_structure: 教学流程（字典）
                - leveled_materials: 分层材料（列表）
                - exercises: 练习题（字典）
                - ppt_outline: PPT大纲（列表）
                - teaching_notes: 教学反思（可选）
            template_vars: 模板变量，包含：
                - teacher_name: 教师姓名
                - school: 学校名称（可选）
                - date: 日期（可选）

        Returns:
            bytes: Word 文档的二进制内容

        Raises:
            ValueError: 如果内容数据无效
            Exception: 如果文档生成失败
        """
        try:
            # 验证必要字段
            if not content.get("title"):
                raise ValueError("教案标题不能为空")

            # 创建新文档
            self.doc = Document()

            # 设置文档默认字体
            self._setup_document_styles()

            # 添加封面页
            self._add_cover_page(content, template_vars)

            # 添加分页
            self.doc.add_page_break()

            # 添加基本信息
            self._add_overview(content, template_vars)

            # 添加教学目标
            if content.get("objectives"):
                self._add_objectives(content["objectives"])

            # 添加教学流程
            if content.get("teaching_structure"):
                self._add_teaching_structure(content["teaching_structure"])

            # 添加核心词汇
            if content.get("vocabulary"):
                self._add_vocabulary_table(content["vocabulary"])

            # 添加语法点
            if content.get("grammar_points"):
                self._add_grammar_points(content["grammar_points"])

            # 添加分层材料
            if content.get("leveled_materials"):
                self._add_leveled_materials(content["leveled_materials"])

            # 添加练习题
            if content.get("exercises"):
                self._add_exercises(content["exercises"])

            # 添加PPT大纲
            if content.get("ppt_outline"):
                self._add_ppt_outline(content["ppt_outline"])

            # 添加教学反思（如果有）
            if content.get("teaching_notes"):
                self.doc.add_page_break()
                self._add_heading("教学反思", level=1)
                p = self.doc.add_paragraph(content["teaching_notes"])

            # 保存到字节流
            doc_bytes = self._save_to_bytes()

            logger.info(f"Word文档生成成功: {content.get('title')}")
            return doc_bytes

        except Exception as e:
            logger.error(f"Word文档生成失败: {str(e)}")
            raise Exception(f"Word文档生成失败: {str(e)}")

    def _setup_document_styles(self) -> None:
        """设置文档默认样式（包括中文字体）"""
        # 设置默认字体
        self.doc.styles["Normal"].font.name = "Calibri"
        self.doc.styles["Normal"].font.size = Pt(self.FONT_SIZE_NORMAL)
        self.doc.styles["Normal"].font.color.rgb = self.COLOR_NORMAL

        # 设置中文字体
        r = self.doc.styles["Normal"].element
        r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def _save_to_bytes(self) -> bytes:
        """
        将文档保存到字节流

        Returns:
            bytes: 文档的二进制内容
        """
        doc_stream = BytesIO()
        self.doc.save(doc_stream)
        doc_stream.seek(0)
        return doc_stream.read()

    def _add_cover_page(self, content: Dict[str, Any], template_vars: Dict[str, Any]) -> None:
        """
        添加封面页

        Args:
            content: 教案内容
            template_vars: 模板变量
        """
        # 标题
        title = content.get("title", "教案")
        h1 = self.doc.add_heading(title, level=0)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_heading_color(h1, self.COLOR_HEADING)

        # 副标题
        subtitle = f"{content.get('level', '')} 等级 - {content.get('topic', '')}"
        h2 = self.doc.add_heading(subtitle, level=2)
        h2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 空行
        self.doc.add_paragraph()

        # 教师信息
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_text_with_style(
            p,
            f"授课教师: {template_vars.get('teacher_name', '')}",
            bold=True
        )

        # 学校信息（如果有）
        if template_vars.get("school"):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_text_with_style(p, f"学校: {template_vars['school']}")

        # 日期信息（如果有）
        if template_vars.get("date"):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_text_with_style(p, f"日期: {template_vars['date']}")

    def _add_overview(self, content: Dict[str, Any], template_vars: Dict[str, Any]) -> None:
        """
        添加课程概述

        Args:
            content: 教案内容
            template_vars: 模板变量
        """
        self._add_heading("基本信息", level=1)

        # 创建概述表格
        table = self.doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        # 添加表格数据
        overview_data = [
            ("教案标题", content.get("title", "")),
            ("课程等级", content.get("level", "")),
            ("课程主题", content.get("topic", "")),
            ("课程时长", f"{content.get('duration', 0)} 分钟"),
        ]

        if content.get("target_exam"):
            overview_data.append(("目标考试", content["target_exam"]))

        overview_data.extend([
            ("授课教师", template_vars.get("teacher_name", "")),
        ])

        if template_vars.get("school"):
            overview_data.append(("学校", template_vars["school"]))

        # 填充表格
        for key, value in overview_data:
            row = table.add_row()
            row.cells[0].text = key
            row.cells[1].text = str(value)

            # 设置表头样式
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(self.FONT_SIZE_NORMAL)

        self.doc.add_paragraph()  # 空行

    def _add_objectives(self, objectives: Dict[str, Any]) -> None:
        """
        添加教学目标

        Args:
            objectives: 教学目标数据
        """
        self._add_heading("教学目标", level=1)

        # 类别名称映射
        category_names = {
            "language_knowledge": "语言知识",
            "language_skills": "语言技能",
            "learning_strategies": "学习策略",
            "cultural_awareness": "文化意识",
            "emotional_attitudes": "情感态度",
        }

        # 技能名称映射
        skill_names = {
            "listening": "听力",
            "speaking": "口语",
            "reading": "阅读",
            "writing": "写作",
        }

        for category, items in objectives.items():
            if not items:
                continue

            # 添加类别标题
            category_name = category_names.get(category, category)
            self._add_heading(category_name, level=2)

            if isinstance(items, dict):
                # 处理嵌套结构（如 language_skills）
                for skill, skill_items in items.items():
                    if skill_items:
                        skill_name = skill_names.get(skill, skill)
                        p = self.doc.add_paragraph(style="List Bullet")
                        self._add_text_with_style(p, f"{skill_name}:", bold=True)

                        for item in skill_items:
                            p = self.doc.add_paragraph(style="List Bullet 2")
                            self._add_text_with_style(p, item)

            elif isinstance(items, list):
                # 处理列表结构
                for item in items:
                    p = self.doc.add_paragraph(style="List Bullet")
                    self._add_text_with_style(p, item)

        self.doc.add_paragraph()  # 空行

    def _add_vocabulary_table(self, vocabulary: Dict[str, Any]) -> None:
        """
        添加词汇表格

        Args:
            vocabulary: 词汇数据，按词性分组
        """
        self._add_heading("核心词汇", level=1)

        # 词性名称映射
        pos_names = {
            "noun": "名词",
            "verb": "动词",
            "adj": "形容词",
            "adv": "副词",
            "prep": "介词",
            "phrase": "短语",
        }

        for pos_type, words in vocabulary.items():
            if not words:
                continue

            # 添加词性标题
            pos_name = pos_names.get(pos_type, pos_type)
            self._add_heading(pos_name, level=2)

            # 创建词汇表格
            table = self.doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"

            # 表头
            headers = ["单词", "音标", "词性", "中文释义", "例句"]
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                # 设置表头样式
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = self.COLOR_HEADING

            # 添加词汇数据
            for word_data in words:
                row = table.add_row()
                row.cells[0].text = word_data.get("word", "")
                row.cells[1].text = word_data.get("phonetic", "")
                row.cells[2].text = word_data.get("part_of_speech", "")
                row.cells[3].text = word_data.get("meaning_cn", "")
                row.cells[4].text = word_data.get("example_sentence", "")

        self.doc.add_paragraph()  # 空行

    def _add_grammar_points(self, grammar_points: List[Dict[str, Any]]) -> None:
        """
        添加语法点

        Args:
            grammar_points: 语法点列表
        """
        self._add_heading("语法点", level=1)

        for gp in grammar_points:
            # 语法点名称
            name = gp.get("name", "")
            if name:
                self._add_heading(name, level=2)

            # 描述
            description = gp.get("description", "")
            if description:
                p = self.doc.add_paragraph()
                self._add_text_with_style(p, "描述: ", bold=True)
                self._add_text_with_style(p, description)

            # 规则
            rule = gp.get("rule", "")
            if rule:
                p = self.doc.add_paragraph()
                self._add_text_with_style(p, "规则: ", bold=True)
                self._add_text_with_style(p, rule)

            # 例句
            examples = gp.get("examples", [])
            if examples:
                p = self.doc.add_paragraph()
                self._add_text_with_style(p, "例句:", bold=True)
                for example in examples:
                    p = self.doc.add_paragraph(style="List Bullet")
                    self._add_text_with_style(p, example)

            # 常见错误
            mistakes = gp.get("common_mistakes", [])
            if mistakes:
                p = self.doc.add_paragraph()
                self._add_text_with_style(p, "常见错误:", bold=True)
                for mistake in mistakes:
                    p = self.doc.add_paragraph(style="List Bullet")
                    self._add_text_with_style(p, mistake, color=self.COLOR_ACCENT)

            # 练习建议
            tips = gp.get("practice_tips", [])
            if tips:
                p = self.doc.add_paragraph()
                self._add_text_with_style(p, "练习建议:", bold=True)
                for tip in tips:
                    p = self.doc.add_paragraph(style="List Bullet")
                    self._add_text_with_style(p, tip)

        self.doc.add_paragraph()  # 空行

    def _add_teaching_structure(self, structure: Dict[str, Any]) -> None:
        """
        添加教学流程

        Args:
            structure: 教学流程数据
        """
        self._add_heading("教学流程", level=1)

        # 阶段名称映射
        phase_names = {
            "warm_up": "热身阶段",
            "presentation": "讲解阶段",
            "practice": "练习阶段",
            "production": "产出阶段",
            "summary": "总结阶段",
            "homework": "课后作业",
        }

        for phase, data in structure.items():
            if not data:
                continue

            # 阶段标题
            phase_name = phase_names.get(phase, phase)
            self._add_heading(phase_name, level=2)

            if isinstance(data, dict):
                # 时长
                duration = data.get("duration", 0)
                if duration:
                    p = self.doc.add_paragraph()
                    self._add_text_with_style(p, f"时长: {duration} 分钟")

                # 描述
                description = data.get("description", "")
                if description:
                    p = self.doc.add_paragraph()
                    self._add_text_with_style(p, "描述: ", bold=True)
                    self._add_text_with_style(p, description)

                # 活动
                activities = data.get("activities", [])
                if activities:
                    p = self.doc.add_paragraph()
                    self._add_text_with_style(p, "活动:", bold=True)
                    for activity in activities:
                        p = self.doc.add_paragraph(style="List Bullet")
                        self._add_text_with_style(p, activity)

                # 教师活动
                teacher_actions = data.get("teacher_actions", [])
                if teacher_actions:
                    p = self.doc.add_paragraph()
                    self._add_text_with_style(p, "教师活动:", bold=True)
                    for action in teacher_actions:
                        p = self.doc.add_paragraph(style="List Bullet")
                        self._add_text_with_style(p, action)

                # 学生活动
                student_actions = data.get("student_actions", [])
                if student_actions:
                    p = self.doc.add_paragraph()
                    self._add_text_with_style(p, "学生活动:", bold=True)
                    for action in student_actions:
                        p = self.doc.add_paragraph(style="List Bullet")
                        self._add_text_with_style(p, action)

                # 所需材料
                materials = data.get("materials", [])
                if materials:
                    p = self.doc.add_paragraph()
                    self._add_text_with_style(p, "所需材料:", bold=True)
                    for material in materials:
                        p = self.doc.add_paragraph(style="List Bullet")
                        self._add_text_with_style(p, material)

        self.doc.add_paragraph()  # 空行

    def _add_leveled_materials(self, materials: List[Dict[str, Any]]) -> None:
        """
        添加分层材料

        Args:
            materials: 分层材料列表
        """
        self._add_heading("分层阅读材料", level=1)

        for material in materials:
            # 材料标题
            title = material.get("title", "")
            level = material.get("level", "")
            if title:
                self._add_heading(f"{title} (CEFR {level})", level=2)

            # 基本信息
            p = self.doc.add_paragraph()
            p.add_run(f"等级: {level} | ")
            p.add_run(f"字数: {material.get('word_count', 0)} 字")

            # 内容
            content = material.get("content", "")
            if content:
                self.doc.add_paragraph(content)

            # 重点词汇
            vocab_list = material.get("vocabulary_list", [])
            if vocab_list:
                p = self.doc.add_paragraph()
                self._add_text_with_style(p, "重点词汇:", bold=True)
                for vocab in vocab_list:
                    p = self.doc.add_paragraph(style="List Bullet")
                    self._add_text_with_style(
                        p,
                        f"{vocab.get('word', '')}: {vocab.get('meaning_cn', '')}"
                    )

            # 理解问题
            questions = material.get("comprehension_questions", [])
            if questions:
                p = self.doc.add_paragraph()
                self._add_text_with_style(p, "理解问题:", bold=True)
                for question in questions:
                    p = self.doc.add_paragraph(style="List Bullet")
                    self._add_text_with_style(p, question)

        self.doc.add_paragraph()  # 空行

    def _add_exercises(self, exercises: Dict[str, Any]) -> None:
        """
        添加练习题

        Args:
            exercises: 练习题数据，按类型分组
        """
        self._add_heading("练习题", level=1)

        # 题型名称映射
        type_names = {
            "multiple_choice": "选择题",
            "fill_blank": "填空题",
            "matching": "匹配题",
            "essay": "写作题",
            "speaking": "口语题",
            "translation": "翻译题",
        }

        for ex_type, exercise_list in exercises.items():
            if not exercise_list:
                continue

            # 题型标题
            type_name = type_names.get(ex_type, ex_type)
            self._add_heading(type_name, level=2)

            for idx, exercise in enumerate(exercise_list, 1):
                # 题目
                question = exercise.get("question", "")
                p = self.doc.add_paragraph()
                self._add_text_with_style(p, f"{idx}. ", bold=True)
                self._add_text_with_style(p, question)

                # 选项（选择题）
                options = exercise.get("options", [])
                if options:
                    option_labels = ["A", "B", "C", "D", "E", "F"]
                    for i, option in enumerate(options):
                        if i < len(option_labels):
                            p = self.doc.add_paragraph(style="List Bullet")
                            self._add_text_with_style(p, f"{option_labels[i]}. {option}")

                # 答案
                answer = exercise.get("correct_answer", "")
                if answer:
                    p = self.doc.add_paragraph()
                    self._add_text_with_style(p, "答案: ", bold=True)
                    self._add_text_with_style(p, str(answer), color=self.COLOR_ACCENT)

                # 解析
                explanation = exercise.get("explanation", "")
                if explanation:
                    p = self.doc.add_paragraph()
                    self._add_text_with_style(p, "解析: ", bold=True)
                    self._add_text_with_style(p, explanation)

                # 空行
                self.doc.add_paragraph()

        self.doc.add_paragraph()  # 空行

    def _add_ppt_outline(self, ppt_outline: List[Dict[str, Any]]) -> None:
        """
        添加 PPT 大纲

        Args:
            ppt_outline: PPT大纲列表
        """
        self._add_heading("PPT大纲", level=1)

        # 创建PPT大纲表格
        table = self.doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        # 表头
        table.rows[0].cells[0].text = "幻灯片"
        table.rows[0].cells[1].text = "内容要点"

        # 设置表头样式
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = self.COLOR_HEADING

        # 添加幻灯片数据
        for slide in ppt_outline:
            row = table.add_row()

            # 幻灯片编号和标题
            slide_num = slide.get("slide_number", "")
            slide_title = slide.get("title", "")
            row.cells[0].text = f"{slide_num}. {slide_title}"

            # 内容要点
            content = slide.get("content", [])
            row.cells[1].text = "\n".join(content) if content else ""

            # 演讲者备注（如果有）
            notes = slide.get("notes", "")
            if notes:
                p = row.cells[1].paragraphs[-1]
                if p.text:
                    p.add_run("\n\n")
                run = p.add_run(f"[备注: {notes}]")
                run.font.color.rgb = RGBColor(128, 128, 128)
                run.font.italic = True

        self.doc.add_paragraph()  # 空行

    # ==================== 辅助方法 ====================

    def _add_heading(self, text: str, level: int = 1) -> None:
        """
        添加标题

        Args:
            text: 标题文本
            level: 标题级别 (1-3)
        """
        heading = self.doc.add_heading(text, level=level)
        self._set_heading_color(heading, self.COLOR_HEADING)

    def _set_heading_color(self, heading, color: RGBColor) -> None:
        """
        设置标题颜色

        Args:
            heading: 标题段落对象
            color: 颜色对象
        """
        for run in heading.runs:
            run.font.color.rgb = color

    def _add_text_with_style(
        self,
        paragraph,
        text: str,
        bold: bool = False,
        italic: bool = False,
        color: Optional[RGBColor] = None
    ) -> None:
        """
        在段落中添加带样式的文本

        Args:
            paragraph: 段落对象
            text: 要添加的文本
            bold: 是否加粗
            italic: 是否斜体
            color: 颜色（可选）
        """
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
