"""
Word文档渲染服务 - AI英语教学系统
支持将Markdown内容转换为Word文档格式
"""
import io
import re
from dataclasses import dataclass
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table


@dataclass
class MarkdownElement:
    """Markdown元素"""

    type: str  # heading, paragraph, list, table, code, quote
    level: int = 1  # 标题级别
    content: str = ""
    children: List["MarkdownElement"] = None


class WordRendererService:
    """
    Word文档渲染服务类

    功能：
    1. 将Markdown内容转换为Word文档
    2. 支持标题、段落、列表、表格、代码块等元素
    3. 中文字体支持
    4. 样式自定义
    """

    # 中文字体映射
    CHINESE_FONTS = {
        "mac": "PingFang SC",
        "windows": "Microsoft YaHei",
        "linux": "WenQuanYi Micro Hei",
        "default": "SimSun",
    }

    # 标题样式映射
    HEADING_STYLES = {
        1: "Title",
        2: "Heading 1",
        3: "Heading 2",
        4: "Heading 3",
        5: "Heading 4",
    }

    def __init__(self):
        """初始化渲染服务"""
        self.current_os = self._detect_os()
        self._doc: Optional[Document] = None
        self._styles_initialized = False

    def _detect_os(self) -> str:
        """检测操作系统"""
        import sys

        platform = sys.platform
        if platform == "darwin":
            return "mac"
        elif platform == "win32":
            return "windows"
        else:
            return "linux"

    def _get_chinese_font(self) -> str:
        """获取中文字体"""
        return self.CHINESE_FONTS.get(self.current_os, self.CHINESE_FONTS["default"])

    def _init_document_styles(self, doc: Document) -> None:
        """初始化文档样式"""
        if self._styles_initialized:
            return

        chinese_font = self._get_chinese_font()

        # 设置默认段落样式
        default_style = doc.styles["Normal"]
        default_style.font.name = chinese_font
        default_style.font.size = Pt(12)
        default_style._element.rPr.rFonts.set(qn("w:eastAsia"), chinese_font)

        # 设置标题样式
        for level in range(1, 5):
            style_name = self.HEADING_STYLES.get(level)
            if style_name and style_name in doc.styles:
                style = doc.styles[style_name]
                style.font.name = chinese_font
                style.font.size = Pt(14 if level == 1 else (12 + (4 - level) * 2))
                style._element.rPr.rFonts.set(qn("w:eastAsia"), chinese_font)

        self._styles_initialized = True

    def parse_markdown(self, markdown: str) -> List[MarkdownElement]:
        """
        解析Markdown为结构化元素

        Args:
            markdown: Markdown文本

        Returns:
            Markdown元素列表
        """
        lines = markdown.split("\n")
        elements: List[MarkdownElement] = []
        i = 0

        while i < len(lines):
            line = lines[i]

            # 跳过空行
            if not line.strip():
                i += 1
                continue

            # 标题 (## 或 #)
            if line.startswith("#"):
                match = re.match(r"^(#{1,6})\s+(.+)$", line)
                if match:
                    level = len(match.group(1))
                    content = match.group(2).strip()
                    elements.append(
                        MarkdownElement(type="heading", level=level, content=content)
                    )
                    i += 1
                    continue

            # 无序列表 (- 或 *)
            if line.strip().startswith(("- ", "* ")):
                list_content = line.strip()[2:]
                # 收集连续的列表项
                list_items = [list_content]
                j = i + 1
                while j < len(lines) and lines[j].strip().startswith(
                    ("- ", "* ", "-", "*")
                ):
                    item = lines[j].strip()[2:] if lines[j].strip()[1] == " " else lines[j].strip()[1:]
                    if item:
                        list_items.append(item)
                    j += 1
                elements.append(MarkdownElement(type="list", content=str(list_items)))
                i = j
                continue

            # 有序列表 (1. 2. 等)
            if re.match(r"^\d+[\.\)]\s+", line):
                list_content = line.strip()
                # 收集连续的有序列表
                list_items = [list_content]
                j = i + 1
                while j < len(lines) and re.match(
                    r"^\d+[\.\)]\s+", lines[j].strip()
                ):
                    list_items.append(lines[j].strip())
                    j += 1
                elements.append(
                    MarkdownElement(type="olist", content=str(list_items))
                )
                i = j
                continue

            # 代码块 (```)
            if line.startswith("```"):
                code_content = []
                j = i + 1
                while j < len(lines) and not lines[j].startswith("```"):
                    code_content.append(lines[j])
                    j += 1
                elements.append(
                    MarkdownElement(
                        type="code", content="\n".join(code_content).strip()
                    )
                )
                i = j + 1 if j < len(lines) else i + 1
                continue

            # 引用 (> )
            if line.startswith("> "):
                quote_content = [line[2:]]
                j = i + 1
                while j < len(lines) and lines[j].startswith("> "):
                    quote_content.append(lines[j][2:])
                    j += 1
                elements.append(
                    MarkdownElement(type="quote", content="\n".join(quote_content))
                )
                i = j
                continue

            # 表格
            if line.startswith("|") and i + 2 < len(lines) and lines[i + 1].startswith("|"):
                table_rows = []
                # 解析表头
                headers = [cell.strip() for cell in line.strip("|").split("|")]
                table_rows.append(headers)
                # 解析分隔行
                separators = [cell.strip() for cell in lines[i + 1].strip("|").split("|")]
                # 解析数据行
                j = i + 2
                while j < len(lines) and lines[j].startswith("|"):
                    row = [cell.strip() for cell in lines[j].strip("|").split("|")]
                    table_rows.append(row)
                    j += 1
                elements.append(MarkdownElement(type="table", content=str(table_rows)))
                i = j
                continue

            # 分割线
            if re.match(r"^[-*_]{3,}$", line.strip()):
                elements.append(MarkdownElement(type="hr"))
                i += 1
                continue

            # 普通段落
            # 处理粗体和斜体
            content = line
            content = re.sub(r"\*\*\*(.+?)\*\*\*", r"\1", content)  # 粗体
            content = re.sub(r"\*\*(.+?)\*\*", r"\1", content)  # 粗体
            content = re.sub(r"__(.+?)__", r"\1", content)  # 下划线
            content = re.sub(r"_(.+?)_", r"\1", content)  # 斜体
            content = re.sub(r"`(.+?)`", r"\1", content)  # 行内代码

            elements.append(MarkdownElement(type="paragraph", content=content))
            i += 1

        return elements

    def render_element(self, element: MarkdownElement, doc: Document) -> None:
        """
        渲染单个Markdown元素到Word文档

        Args:
            element: Markdown元素
            doc: Word文档对象
        """
        chinese_font = self._get_chinese_font()

        if element.type == "heading":
            # 标题
            style_name = self.HEADING_STYLES.get(element.level, "Heading 1")
            para = doc.add_paragraph(style=style_name)
            para.add_run(element.content)

        elif element.type == "paragraph":
            # 普通段落
            doc.add_paragraph(element.content)

        elif element.type == "list":
            # 无序列表
            try:
                items = eval(element.content) if element.content.startswith("[") else [element.content]
                for item in items:
                    # 移除列表标记
                    item = re.sub(r"^[-*]\s*", "", str(item)).strip()
                    para = doc.add_paragraph()
                    para.style = doc.styles["List Bullet"]
                    para.add_run(item)
            except Exception:
                doc.add_paragraph(element.content)

        elif element.type == "olist":
            # 有序列表
            try:
                items = eval(element.content) if element.content.startswith("[") else [element.content]
                for idx, item in enumerate(items, 1):
                    # 移除列表标记
                    item = re.sub(r"^\d+[\.\)]\s*", "", str(item)).strip()
                    para = doc.add_paragraph()
                    para.style = doc.styles["List Number"]
                    para.add_run(item)
            except Exception:
                doc.add_paragraph(element.content)

        elif element.type == "code":
            # 代码块
            para = doc.add_paragraph()
            run = para.add_run(element.content)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)

        elif element.type == "quote":
            # 引用块
            para = doc.add_paragraph()
            run = para.add_run(f"「{element.content}」")
            run.font.italic = True

        elif element.type == "table":
            # 表格
            try:
                rows = eval(element.content) if element.content.startswith("[") else [[element.content]]
                if rows:
                    num_cols = len(rows[0])
                    table = doc.add_table(rows=len(rows), cols=num_cols)
                    table.style = "Table Grid"

                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_text in enumerate(row_data):
                            cell = table.cell(row_idx, col_idx)
                            cell.text = str(cell_text)
                            if row_idx == 0:  # 表头
                                cell.paragraphs[0].runs[0].bold = True
            except Exception:
                pass

        elif element.type == "hr":
            # 分割线
            doc.add_paragraph("_" * 30)

    def markdown_to_docx_bytes(
        self,
        markdown: str,
        title: str = "",
        author: str = "AI英语教学系统",
    ) -> bytes:
        """
        将Markdown转换为Word文档（返回字节）

        Args:
            markdown: Markdown文本
            title: 文档标题
            author: 作者

        Returns:
            Word文档字节数据
        """
        # 创建文档
        doc = Document()
        self._doc = doc
        self._init_document_styles(doc)

        # 设置文档属性
        doc.core_properties.title = title
        doc.core_properties.author = author
        doc.core_properties.created = datetime.now()
        doc.core_properties.modified = datetime.now()

        # 解析并渲染Markdown
        elements = self.parse_markdown(markdown)

        # 标题
        if title:
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.font.size = Pt(18)
            title_run.font.bold = True
            title_run.font.name = self._get_chinese_font()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_after = Pt(18)

        # 渲染所有元素
        for element in elements:
            self.render_element(element, doc)

        # 保存到字节流
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()

    def markdown_to_docx_file(
        self,
        markdown: str,
        output_path: str,
        title: str = "",
        author: str = "AI英语教学系统",
    ) -> None:
        """
        将Markdown保存为Word文件

        Args:
            markdown: Markdown文本
            output_path: 输出文件路径
            title: 文档标题
            author: 作者
        """
        docx_bytes = self.markdown_to_docx_bytes(markdown, title, author)

        with open(output_path, "wb") as f:
            f.write(docx_bytes)

    def add_toc(self, doc: Document, title: str = "目录") -> None:
        """
        添加目录（占位实现）

        Args:
            doc: Word文档对象
            title: 目录标题
        """
        # Word原生目录需要复杂处理，这里添加占位符
        para = doc.add_paragraph()
        para.add_run(title).bold = True
        para.add_run("\n（请在Word中更新目录域）")


# 创建服务工厂函数
def get_word_renderer_service() -> WordRendererService:
    """
    获取Word渲染服务实例

    Returns:
        WordRendererService: Word渲染服务实例
    """
    return WordRendererService()
